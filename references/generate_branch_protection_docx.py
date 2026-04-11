#!/usr/bin/env python3
"""Generate branded branch-protection-guide.docx using nuDesk brand colors."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
PRIMARY_DARK = RGBColor(0x1A, 0x1A, 0x2E)
PRIMARY_ACCENT = RGBColor(0x25, 0x63, 0xEB)  # nuDesk Blue
SECONDARY_ACCENT = RGBColor(0x0D, 0x94, 0x88)  # nuDesk Teal
SUPPORTING_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = RGBColor(0xF9, 0xFA, 0xFB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "branch-protection-guide.docx")

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# --- Footer ---
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("nuDesk LLC | AI-Workforce Agency for Financial Services")
run.font.name = "Calibri"
run.font.size = Pt(9)
run.font.color.rgb = SUPPORTING_GRAY


# --- Helper functions ---
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level <= 2:
            run.font.color.rgb = PRIMARY_DARK
            run.bold = True
        else:
            run.font.color.rgb = PRIMARY_DARK
            run.bold = True
    # Remove spacing after for tighter layout
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.space_before = Pt(12) if level > 1 else Pt(18)
    return h


def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_rich(segments):
    """Add a paragraph with mixed formatting. segments = [(text, bold, italic, color), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic, color in segments:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(lines):
    """Add a code block with gray background."""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = PRIMARY_DARK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Light gray shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_table(headers, rows):
    """Add a branded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True
        # Dark background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A2E" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle code formatting (backticks)
            clean = val.strip("`") if val.startswith("`") and val.endswith("`") else val
            run = p.add_run(clean)
            run.font.name = "Consolas" if val.startswith("`") else "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY_DARK
            # Alternating row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FAFB" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table


def add_callout(text, border_color=PRIMARY_ACCENT):
    """Add a callout box with left border accent."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = SUPPORTING_GRAY
    # Left border via paragraph border
    pPr = p.paragraph_format.element.get_or_add_pPr()
    hex_color = f"{border_color[0]:02X}{border_color[1]:02X}{border_color[2]:02X}"
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{hex_color}"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_after = Pt(8)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = "Calibri"
        run_b.font.size = Pt(11)
        run_b.bold = True
        run_b.font.color.rgb = PRIMARY_DARK
        run_rest = p.add_run(text)
        run_rest.font.name = "Calibri"
        run_rest.font.size = Pt(11)
    else:
        # Clear default and re-add with formatting
        p.clear()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_divider():
    """Add a thin blue horizontal divider."""
    p = doc.add_paragraph()
    pPr = p.paragraph_format.element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2563EB"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)


# ============================================================
# BUILD THE DOCUMENT
# ============================================================

# Title
title = doc.add_heading("Branch Protection & PR Workflow", level=0)
for run in title.runs:
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY_DARK
    run.font.size = Pt(26)

# Subtitle line
add_body_rich([
    ("nuDesk LLC", True, False, PRIMARY_ACCENT),
])

# Metadata callout
add_callout("Effective: 2026-04-10  |  SOC 2 Control: SD-02  |  Applies to: All nuDesk GitHub repos")

add_divider()

# --- WHY ---
add_heading_styled("Why", level=1)
add_body("SOC 2 requires every production change to go through a documented review process. "
         "Branch protection enforces this by requiring a pull request before code reaches main.")
add_body("Every PR creates an audit trail \u2014 who changed what, when, and why.")

add_divider()

# --- THE 30-SECOND VERSION ---
add_heading_styled("The 30-Second Version", level=1)

add_body_rich([
    ("Before (old way):", True, False, PRIMARY_DARK),
])
add_code_block([
    'git add . && git commit -m "feat: thing" && git push origin main',
])

doc.add_paragraph()  # spacing

add_body_rich([
    ("After (new way):", True, False, PRIMARY_DARK),
])
add_code_block([
    "git checkout -b feat/thing",
    'git add . && git commit -m "feat: thing"',
    "git push -u origin feat/thing",
    'gh pr create --title "feat: thing" --body "Description here"',
    "gh pr merge --squash --delete-branch",
])

doc.add_paragraph()

add_body_rich([
    ("Or just ask Claude:", True, False, PRIMARY_ACCENT),
])
add_callout('"Commit these changes and create a PR"\n\nClaude Code handles branching, pushing, and PR creation automatically.')

add_divider()

# --- WORKFLOW BY SCENARIO ---
add_heading_styled("Workflow by Scenario", level=1)

# Scenario 1
add_heading_styled("Scenario 1: Quick Fix (5 min)", level=2)
add_code_block([
    "git checkout -b fix/typo-in-readme",
    "# make your fix",
    "git add README.md",
    'git commit -m "fix: correct typo in README"',
    "git push -u origin fix/typo-in-readme",
    "gh pr create --fill    # auto-fills title+body from commit",
    "gh pr merge --squash --delete-branch",
])

doc.add_paragraph()

# Scenario 2
add_heading_styled("Scenario 2: Feature Work (Multi-Commit)", level=2)
add_code_block([
    "git checkout -b feat/new-skill",
    "# work, commit, work, commit...",
    "git push -u origin feat/new-skill",
    'gh pr create --title "feat: add meeting-prep skill" --body "Adds the meeting-prep skill with..."',
    "# Review the PR yourself or wait for the other person",
    "gh pr merge --squash --delete-branch",
])

doc.add_paragraph()

# Scenario 3
add_heading_styled("Scenario 3: Using Claude Code (Easiest)", level=2)
add_body("Just tell Claude what you want. Examples:")
add_bullet('"Make this change and create a PR"')
add_bullet('"Commit everything and open a PR with a good description"')
add_bullet('/commit \u2014 Claude\'s built-in commit skill (creates PR if on protected branch)')

# Scenario 4
add_heading_styled("Scenario 4: Overnight Agent Queue", level=2)
add_body("The agent creates a feature branch and opens a PR automatically. "
         "Sean reviews and merges the next morning.")

# Scenario 5
add_heading_styled("Scenario 5: Emergency (Bypass)", level=2)
add_body("Admins can push directly to main. Every bypass is logged in GitHub\u2019s audit log. "
         "Use sparingly \u2014 the audit trail is permanent.")
add_code_block([
    "git push origin main  # works for admins, logged as bypass",
])

add_divider()

# --- BRANCH NAMING ---
add_heading_styled("Branch Naming", level=1)

add_table(
    ["Prefix", "Use For", "Example"],
    [
        ["`feat/`", "New features or skills", "`feat/meeting-prep`"],
        ["`fix/`", "Bug fixes", "`fix/version-sync`"],
        ["`chore/`", "Maintenance, cleanup", "`chore/update-deps`"],
        ["`docs/`", "Documentation only", "`docs/add-brand-guide`"],
        ["`agent-queue/`", "Overnight agent work", "`agent-queue/2026-04-10`"],
    ],
)

add_divider()

# --- AFTER MERGING ---
add_heading_styled("After Merging", level=1)

add_body("For the nudesk-os-plugin repo specifically:")
add_code_block([
    "claude plugin update nudesk-os@marketplace",
])

doc.add_paragraph()
add_body("For other repos \u2014 no extra step needed. The merge to main IS the deploy trigger "
         "(or deploy manually as before).")

add_divider()

# --- EASIEST WAY TO STAY COMPLIANT ---
add_heading_styled("Easiest Way to Stay Compliant", level=1)

add_body("Run /session-closeout at the end of every Claude Code session. It automatically:")
add_bullet("Checks for uncommitted or unpushed changes")
add_bullet("Creates a feature branch and PR if work is outstanding")
add_bullet("Self-merges the PR immediately (no approval needed)")
add_bullet("Returns you to main in a clean state")

doc.add_paragraph()
add_callout("This is the safety net \u2014 even if you forget to create a branch during the session, "
            "the closeout catches it and routes everything through a PR before you close the terminal.")

add_divider()

# --- FAQ ---
add_heading_styled("FAQ", level=1)

faqs = [
    ("Q: Can I self-merge my own PR?",
     "A: Yes. Zero approvals required. The PR exists for the audit trail, not as a blocker."),
    ("Q: What if the other person is asleep/unavailable?",
     "A: Self-merge. The PR + your name on it is the audit evidence."),
    ("Q: Do I need to do this for every tiny change?",
     "A: Yes \u2014 but it\u2019s ~30 seconds extra. Claude Code makes it even faster."),
    ("Q: What about repos I rarely touch?",
     "A: Same rule applies. The org ruleset covers all repos automatically."),
    ("Q: Can I still use git commit locally without a branch?",
     "A: Yes \u2014 commits are local. The protection only kicks in when you push to main. Just create the branch before pushing."),
]

for question, answer in faqs:
    add_body(question, bold=True, color=PRIMARY_DARK)
    add_body(answer)

# --- Save ---
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
